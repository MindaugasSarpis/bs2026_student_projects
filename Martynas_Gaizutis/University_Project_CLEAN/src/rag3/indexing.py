from __future__ import annotations

import logging
from pathlib import Path
from typing import Iterable
from uuid import uuid4

import pandas as pd
from docx import Document as DocxDocument
from haystack import Document
from haystack.components.embedders import SentenceTransformersDocumentEmbedder
from pypdf import PdfReader

from rag3.config import AppConfig

logger = logging.getLogger(__name__)


class DocumentIndexer:
    def __init__(self, config: AppConfig, document_store):
        self.config = config
        self.document_store = document_store
        self.index_name = config.elasticsearch.index

    def clear_index(self) -> None:
        logger.info("Clearing index '%s'", self.index_name)
        self.document_store.client.delete_by_query(
            index=self.index_name,
            body={"query": {"match_all": {}}},
            refresh=True,
            conflicts="proceed",
        )

    def index_documents(
        self,
        txt_folder_path: Path | None = None,
        pdf_folder_path: Path | None = None,
        word_folder_path: Path | None = None,
        excel_folder_path: Path | None = None,
        *,
        clear: bool = False,
    ) -> int:
        paths = self._resolve_paths(
            txt_folder_path,
            pdf_folder_path,
            word_folder_path,
            excel_folder_path,
        )

        if clear:
            self.clear_index()

        docs = self._load_documents(paths)
        if not docs:
            logger.warning("No documents found for indexing.")
            return 0

        embedder = SentenceTransformersDocumentEmbedder(model=self.config.embedder.model)
        if self.config.embedder.warm_up:
            embedder.warm_up()

        embedded_docs = embedder.run(docs)["documents"]
        self.document_store.write_documents(embedded_docs)

        indexed_count = self.document_store.client.count(index=self.index_name)["count"]
        logger.info("Index now contains %s documents", indexed_count)
        return len(embedded_docs)

    def _resolve_paths(
        self,
        txt_folder_path: Path | None,
        pdf_folder_path: Path | None,
        word_folder_path: Path | None,
        excel_folder_path: Path | None,
    ) -> dict[str, Path]:
        defaults = self.config.index_documents
        return {
            "txt": Path(txt_folder_path or defaults.txt_folder_path),
            "pdf": Path(pdf_folder_path or defaults.pdf_folder_path),
            "word": Path(word_folder_path or defaults.word_folder_path),
            "excel": Path(excel_folder_path or defaults.excel_folder_path),
        }

    def _load_documents(self, paths: dict[str, Path]) -> list[Document]:
        docs: list[Document] = []
        docs.extend(self._load_txt_files(paths["txt"]))
        docs.extend(self._load_pdf_files(paths["pdf"]))
        docs.extend(self._load_word_files(paths["word"]))
        docs.extend(self._load_excel_files(paths["excel"]))
        logger.info("Loaded %s raw documents", len(docs))
        return docs

    def _load_txt_files(self, folder: Path) -> Iterable[Document]:
        return self._load_text_files(folder, "*.txt", self._read_text_file, "txt")

    def _load_pdf_files(self, folder: Path) -> Iterable[Document]:
        return self._load_text_files(folder, "*.pdf", self._read_pdf_file, "pdf")

    def _load_word_files(self, folder: Path) -> Iterable[Document]:
        return self._load_text_files(folder, "*.docx", self._read_docx_file, "docx")

    def _load_excel_files(self, folder: Path) -> Iterable[Document]:
        return self._load_text_files(folder, "*.xlsx", self._read_excel_file, "xlsx")

    def _load_text_files(self, folder: Path, pattern: str, reader, source_type: str) -> list[Document]:
        if not folder.exists():
            logger.warning("Skipping missing folder: %s", folder)
            return []

        documents: list[Document] = []
        for file_path in sorted(folder.glob(pattern)):
            try:
                content = reader(file_path)
                if not content.strip():
                    continue
                documents.append(
                    Document(
                        content=content,
                        meta={
                            "id": str(uuid4()),
                            "source_path": str(file_path),
                            "source_type": source_type,
                        },
                    )
                )
            except Exception as exc:
                logger.exception("Failed reading %s: %s", file_path, exc)

        return documents

    @staticmethod
    def _read_text_file(file_path: Path) -> str:
        return file_path.read_text(encoding="utf-8", errors="ignore")

    @staticmethod
    def _read_pdf_file(file_path: Path) -> str:
        reader = PdfReader(str(file_path))
        return "\n".join((page.extract_text() or "") for page in reader.pages)

    @staticmethod
    def _read_docx_file(file_path: Path) -> str:
        doc = DocxDocument(file_path)
        return "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())

    @staticmethod
    def _read_excel_file(file_path: Path) -> str:
        data = pd.read_excel(file_path)
        return data.to_string(index=False)
